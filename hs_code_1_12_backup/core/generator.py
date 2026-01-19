"""
[Report Generator]
문서 분석 및 보고서 자동 생성 엔진 (Brain)
사용자 요청을 분석하여 1) 구조화 -> 2) 툴 플래닝 -> 3) 실행 -> 4) 작성을 수행하는 4단계 파이프라인
"""

import json
import logging
import re
from datetime import datetime
from typing import Dict, List, Any, Optional
from pathlib import Path

# DOCX 생성을 위한 라이브러리 (조건부 임포트)
try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    Document = None

# 모듈 의존성
from core.agent_manager import exaone_agent
from core.chat_manager import ChatManager
from services.rag_engine import AdvancedRAGSystem

logger = logging.getLogger(__name__)

class ReportGenerator:
    """
    [핵심 클래스] 보고서 생성 파이프라인
    
    프로세스 개요:
    1. parse_document: 텍스트를 JSON 구조(날짜, 지역, 섹션들)로 변환
    2. plan_tools: 각 섹션을 채우기 위해 어떤 데이터(기상청, 검색 등)가 필요한지 판단
    3. fill_report: 계획된 툴을 실행(API 호출)하고, RAG로 규정을 찾아 내용을 채움
    4. export_to_docx: 최종 결과를 워드 파일로 저장
    """

    def __init__(self):
        # LLM 클라이언트 및 Agent 초기화
        self._chat_manager = ChatManager()
        self.client = self._chat_manager.exaone_client
        self.agent = exaone_agent
        
        # 규정 검색을 위한 RAG 엔진
        self.rag = AdvancedRAGSystem()
        
        # 결과물 저장 경로 설정
        self.output_dir = Path("hs_code/data/output")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.template_path = Path("hs_code/data/template.docx")

    def generate_daily_report_file(self) -> tuple[str, str]:
        """(UI 호환용)"""
        return "report.md", "Logic executed via chat interface."

    def parse_document(self, text: str) -> Dict[str, Any]:
        """
        [Phase 1] 문서 구조 분석 (Structure Analysis)
        
        - 목적: 비정형 텍스트(템플릿, 요청사항)에서 메타데이터와 섹션 구조를 추출
        - 방식: System Prompt를 통해 엄격한 JSON 스키마를 따르도록 유도
        """
        logger.info("📄 [PHASE 1] 문서 구조화 중...")

        # [프롬프트 설계] JSON 포맷 강제 및 누락 정보 확인 요청 (예시 강화)
        system_prompt = """You are a specialized Document Analyst.
Your goal is to parse the input text into a structured JSON execution plan for a disaster report.

RESPONSE FORMAT:
You must output ONLY a valid JSON object. Do not add any explanation.

JSON SCHEMA EXAMPLES:

Case 1: Missing Information (Date or Region)
{
  "status": "incomplete",
  "missing_fields": ["Date", "Region"],
  "clarification_question": "보고서 작성을 위해 [날짜]와 [지역] 정보가 필요합니다."
}

Case 2: Complete Information
{
  "status": "complete",
  "document_type": "Disaster Report",
  "report_date": "2024-07-15",
  "target_region": "용인시",
  "sections": [
    {
      "section_id": "1",
      "title": "기상 전망",
      "original_text": "오늘 용인시 날씨는...",
      "search_query": "풍수해 위기관리 매뉴얼 기상 전망",
      "requires_tools": true
    },
    {
      "section_id": "2",
      "title": "조치 사항",
      "original_text": "비상근무를 실시한다.",
      "search_query": "풍수해 비상근무 기준",
      "requires_tools": false
    }
  ]
}
"""
        
        user_prompt = f"""
[Input Text]
{text[:3000]}

[Instruction]
Extract metadata (Date, Region) and sections from the text above.
If Date or Region is missing, return status 'incomplete'.
Output VALID JSON ONLY.
"""
        
        # [Retry 로직] LLM이 깨진 JSON을 줄 경우 최대 3번 재시도
        for attempt in range(3):
            try:
                logger.info(f"🔄 [PHASE 1] Attempt {attempt+1}/3: LLM 호출...")
                response = self.client.generate_response(
                    messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
                    enable_thinking=False, # 구조 분석은 추론 불필요 (빠른 응답)
                    temperature=0.1, stream=False
                )
                
                if not response or not response.choices: continue
                message = response.choices[0].message
                content = message.content
                if not content: continue
                
                # [JSON 파싱 로직 강화]
                # 1. Markdown 코드 블록 제거
                clean_content = re.sub(r'```json\s*', '', content)
                clean_content = re.sub(r'```\s*', '', clean_content)
                
                # 2. 가장 바깥쪽 중괄호 찾기 (탐욕적 매칭 방지)
                # 단순 {.*} 대신 중괄호 균형을 맞추거나, 가장 넓은 범위를 잡음
                json_match = re.search(r'(\{.*\})', clean_content, re.DOTALL)
                
                if json_match:
                    json_str = json_match.group(1)
                    try:
                        structure = json.loads(json_str)
                        # 필수 정보 누락 시 사용자에게 질문 반환
                        if structure.get("status") == "incomplete": return structure
                        
                        # 섹션 검증
                        if not structure.get("sections"):
                            logger.warning("⚠️ 파싱된 JSON에 'sections'가 비어있습니다.")
                            continue

                        logger.info(f"✅ [PHASE 1] 구조 분석 완료: {len(structure.get('sections', []))}개 섹션")
                        return structure
                    except json.JSONDecodeError as e:
                        logger.error(f"❌ JSON Decode Error: {e}\nContent: {json_str[:200]}...")
                else:
                    logger.error("❌ JSON 패턴을 찾을 수 없습니다.")

            except Exception as e: logger.error(f"Error: {e}")
            
        return {"status": "error", "sections": []}

    def plan_tools(self, structure: Dict[str, Any]) -> Dict[str, Any]:
        """
        [Phase 2] 도구 사용 계획 수립 (Tool Planning)
        
        - 목적: 각 섹션 내용을 채우기 위해 '어떤 툴'이 필요한지 결정
        - 방식: 섹션의 문맥(Context)을 보고 가용 툴(Available Tools) 목록에서 선택
        """
        logger.info("🧠 [PHASE 2] 툴 플래닝 중...")
        
        sections = structure.get("sections", [])
        report_date = structure.get("report_date", "")
        target_region = structure.get("target_region", "")

        # 가용 툴 목록을 프롬프트에 명시하여 Hallucination 방지
        system_prompt = """You are a Strategic Planning Agent.
Decide which tools are needed to fetch real data for the report section.

Available Tools:
- kma_get_ultra_srt_ncst: Current Weather (Temp, Rain, Wind)
- kma_get_vilage_fcst: Short-term Forecast
- kma_get_mid_land_fcst: Mid-term Forecast
- kma_get_wthr_wrn_msg: Weather Alerts (Warnings)
- kma_get_eqk_msg_list: Earthquake Info
- serpapi_web_search: News, Events

Output Format (JSON List): [{"tool": "tool_name", "arguments": {"arg": "val"}}]
If no tools needed, output [].
"""
        
        for section in sections:
            # 도구가 필요 없다고 마킹된 섹션은 패스
            if not section.get("requires_tools", False):
                section["tool_plan"] = []
                continue
            
            user_prompt = f"""
[Context] Date: {report_date}, Region: {target_region}
[Section] {section.get("title")}
[Content] {section.get("original_text")}

Which tools do I need?
"""
            try:
                # 툴 플래닝은 정확도가 생명이므로 Temperature를 낮춤
                response = self.client.generate_response(
                    messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
                    enable_thinking=False, temperature=0.1, stream=False
                )
                
                content = response.choices[0].message.content
                if not content: content = ""
                
                # JSON 리스트([]) 추출
                json_match = re.search(r'\[.*\]', content, re.DOTALL)
                tool_plan = json.loads(json_match.group(0)) if json_match else []
                
                section["tool_plan"] = tool_plan
                logger.info(f"  📌 '{section.get('title')}' 계획: {len(tool_plan)}개 툴")
                
            except Exception as e:
                logger.warning(f"  ⚠️ 플래닝 실패 ({section.get('title')}): {e}")
                section["tool_plan"] = []
        
        return structure

    def fill_report(self, structure: Dict[str, Any]) -> Dict[str, Any]:
        """
        [Phase 3 & 4] 실행 및 작성 (Execution & Writing)
        
        - 목적: 실제 데이터를 가져오고(Action), 최종 문장(Report)을 생성
        - 프로세스:
          1. RAG 검색: 해당 섹션과 관련된 규정/매뉴얼 검색
          2. Tool 실행: Phase 2에서 계획된 툴을 실행하여 실시간 데이터 획득
          3. LLM 작성: [규정 + 실시간 데이터 + 원본 템플릿]을 조합하여 문장 생성
        """
        logger.info("🛠️ [PHASE 3] 실행 및 내용 작성 중...")
        
        # Markdown 보고서 헤더 작성
        report_md = f"# {structure.get('document_type', 'Daily Report')}\n"
        report_md += f"**일시:** {structure.get('report_date')} {structure.get('report_time', '')}\n"
        report_md += f"**지역:** {structure.get('target_region')}\n\n---\n"

        meta_date = structure.get("report_date", "")
        meta_region = structure.get("target_region", "")

        from services.agent_tools import exaone_agent_tools

        for section in structure.get("sections", []):
            title = section.get("title", "Untitled")
            sec_id = section.get("section_id", "")
            original_text = section.get("original_text", "")
            search_query = section.get("search_query", title)
            requires_tools = section.get("requires_tools", False)
            tool_plan = section.get("tool_plan", [])
            
            report_md += f"## {sec_id}. {title}\n\n"
            
            # 툴 사용이 필요 없으면 원본 텍스트 유지
            if not requires_tools and not tool_plan:
                generated_text = original_text
                section["generated_content"] = generated_text
                report_md += f"{generated_text}\n\n"
                continue

            # 1. [RAG Retrieval] 규정/매뉴얼 검색
            rag_docs = self.rag.search(search_query, top_k=2)
            manual_context = ""
            rag_evidence = []
            if rag_docs:
                manual_context = "\n".join([f"- [출처: {Path(d['metadata']['source']).name}] {d['text'][:500]}..." for d in rag_docs])
                rag_evidence = [f"📄 **[규정]** {Path(d['metadata']['source']).name} (관련도 {d['similarity_score']:.2f})" for d in rag_docs]
            else:
                manual_context = "(관련 매뉴얼을 찾을 수 없음)"

            # 2. [Tool Execution] 계획된 툴 실행
            tool_outputs = []
            tool_logs = []
            
            if tool_plan:
                logger.info(f"⚡ [EXEC] '{title}' 섹션용 툴 {len(tool_plan)}개 실행...")
                for plan in tool_plan:
                    t_name = plan.get("tool")
                    t_args = plan.get("arguments", {})
                    
                    friendly_name = "기상청 API" if "kma" in t_name else "웹 검색"
                    
                    try:
                        # 위치 정보가 누락된 경우 메타데이터에서 자동 주입 (Context Injection)
                        if t_name.startswith("kma_") and "location" not in t_args:
                            t_args["location"] = meta_region
                        
                        # 실제 실행 (agent_tools.py 호출)
                        result = exaone_agent_tools.execute_tool(t_name, t_args)
                        
                        # 결과 저장
                        tool_outputs.append(f"[출처: {friendly_name} ({t_name})]\nData: {result}")
                        tool_logs.append(f"- 🔧 **{friendly_name}**: `{str(result)[:100]}...`")
                        
                    except Exception as e:
                        logger.error(f"Tool execution failed: {e}")
                        tool_outputs.append(f"[Tool: {t_name}] Error: {str(e)}")

            # 3. [Content Generation] 최종 문장 작성
            live_data_context = "\n\n".join(tool_outputs) if tool_outputs else "(No live data collected)"
            
            prompt = f"""
[Role]
You are a public official writing a disaster report section.

[Context]
Date: {meta_date}
Region: {meta_region}

[Reference Manual / Regulations]
{manual_context}

[Live Data (Collected from Tools)]
{live_data_context}

[Original Template]
{original_text}

[Directives]
1. Based on the [Live Data], rewrite the "{title}" section.
2. Compare data with [Reference Manual] criteria.
3. If data is missing or normal, state "특이사항 없음" or similar.
4. Use official terminology.
5. Maintain the format of the Original Template.
6. [IMPORTANT] **Transparency & Citation**:
   - For every factual statement, append the source in brackets at the end of the sentence.
   - Example: "현재 용인시 기온은 24도입니다. (출처: 기상청 API)"
   - Do NOT invent sources. Only use provided contexts.
"""
            
            try:
                response = self.client.generate_response(
                    messages=[{"role": "user", "content": prompt}],
                    enable_thinking=True, # 고품질 작성을 위해 Thinking 활성화
                    temperature=0.2, stream=False
                )
                
                content = response.choices[0].message.content
                # Thinking 태그(<think>...</think>) 제거 후 본문만 추출
                final_text = re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL).strip()
                if not final_text: final_text = content
                
                section["generated_content"] = final_text
                report_md += f"{final_text}\n"
                
                # 근거 자료(규정, 로그)를 Markdown 주석으로 첨부
                if rag_evidence or tool_logs:
                    report_md += "\n"
                    if rag_evidence: report_md += "\n".join(rag_evidence) + "\n"
                    if tool_logs: report_md += "\n".join(tool_logs) + "\n\n"
                else:
                    report_md += "\n"

            except Exception as e:
                logger.error(f"❌ Generation failed for {title}: {e}")
                report_md += f"⚠️ Section Error: {str(e)}\n\n"

        structure["full_report_md"] = report_md
        return structure

    def export_to_docx(self, structure: Dict[str, Any]) -> str:
        """
        [Export] DOCX 파일 생성
        - 템플릿(template.docx)에 내용을 채워 넣거나 새 문서 생성
        - 테이블/문단 위치를 찾아 텍스트 치환
        """
        # ... (DOCX 생성 로직 생략, 기존 코드 유지) ...
        if not Document:
            logger.error("python-docx module not found.")
            return ""

        try:
            logger.info("💾 [DOCX] Exporting to DOCX...")
            
            # 템플릿 로드 또는 새 문서 생성
            if self.template_path.exists():
                doc = Document(self.template_path)
            else:
                doc = Document()
                doc.add_heading(structure.get('document_type', 'Disaster Report'), 0)
            
            # 메타데이터 채우기 (문서 상단)
            report_date = structure.get('report_date', '')
            target_region = structure.get('target_region', '')
            
            # 섹션별 내용 채우기
            for section in structure.get("sections", []):
                title = section.get("title", "")
                content = section.get("generated_content", "")
                
                # 1. 표(Table) 검색
                inserted = False
                for table in doc.tables:
                    for row in table.rows:
                        # 첫 번째 셀이 제목과 유사한지 확인
                        if row.cells and title.replace(" ", "") in row.cells[0].text.replace(" ", ""):
                            # 두 번째 셀에 내용 삽입 (있다면)
                            if len(row.cells) > 1:
                                row.cells[1].text = content
                                inserted = True
                                break
                    if inserted: break
                
                # 2. 문단(Paragraph) 검색 (표에 없으면 본문 검색)
                if not inserted:
                    for i, para in enumerate(doc.paragraphs):
                        if title.replace(" ", "") in para.text.replace(" ", ""):
                            # 제목 문단을 찾음. 그 아래 문단에 내용 추가
                            if i + 1 < len(doc.paragraphs):
                                next_para = doc.paragraphs[i+1]
                                if not next_para.text.strip():
                                    next_para.text = content
                                else:
                                    para.insert_paragraph_before(content)
                            else:
                                doc.add_paragraph(content)
                            inserted = True
                            break
                
                # 3. 못 찾았으면 문서 끝에 추가
                if not inserted:
                    doc.add_heading(title, level=1)
                    doc.add_paragraph(content)
            
            # 파일 저장
            filename = f"Report_{{target_region}}_{{report_date}}.docx".replace(" ", "_")
            output_path = self.output_dir / filename
            doc.save(output_path)
            
            logger.info(f"✅ DOCX saved: {output_path}")
            return str(output_path)

        except Exception as e:
            logger.error(f"❌ DOCX Export failed: {e}")
            return ""